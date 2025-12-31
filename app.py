import streamlit as st
import google.generativeai as genai
from docx import Document
from io import BytesIO

# --- 0. 全局設定 (請在此修改您的商品資訊) ---
PAGE_TITLE = "企業級 SOP 智能生成器 "
GUMROAD_LINK = "https://louisian5723.gumroad.com/l/wjxao"  # 👈 請記得換成您真的 Gumroad 商品連結
MAX_USAGE_PER_SESSION = 10  # 限制每次登入只能生成幾次 (防止惡意刷單)

# --- 1. 頁面初始化 ---
st.set_page_config(page_title=PAGE_TITLE, page_icon="📝", layout="wide")

# 初始化 Session State (用於計數器和儲存結果)
if 'usage_count' not in st.session_state:
    st.session_state['usage_count'] = 0
if 'result' not in st.session_state:
    st.session_state['result'] = ""

# --- 2. 核心功能函數 ---

def check_password():
    """檢查用戶密碼是否正確，錯誤則顯示購買連結"""
    # 如果雲端沒設密碼，為了不報錯，暫時放行 (或您可以選擇報錯)
    if "APP_PASSWORD" not in st.secrets:
        return True
    
    st.sidebar.header("🔐 會員登入")
    password_input = st.sidebar.text_input("請輸入通行密碼", type="password")
    
    if password_input == st.secrets["APP_PASSWORD"]:
        return True
    else:
        st.sidebar.divider()
        st.sidebar.warning("🔒 請輸入密碼以解鎖功能")
        st.sidebar.markdown(f"""
        ### 尚未擁有密碼？
        本工具為付費軟體，提供企業級 SOP 生成服務。
        
        👉 **[點擊購買 VIP 通行證 (US$ 5)]({GUMROAD_LINK})**
        
        *付款後，系統將自動發送密碼至您的信箱。*
        """)
        return False

def check_usage_limit():
    """檢查是否超過使用次數"""
    if st.session_state['usage_count'] >= MAX_USAGE_PER_SESSION:
        st.error(f"🚫 您已達到本次登入的使用上限 ({MAX_USAGE_PER_SESSION} 次)。為了確保服務品質，請稍後重新整理頁面再試。")
        return False
    return True

def generate_sop(raw_text):
    """呼叫 Gemini 生成 SOP"""
    # 檢查 API Key
    if "GEMINI_API_KEY" not in st.secrets:
        st.error("❌ 系統錯誤：未設定 API Key。")
        return None
        
    api_key = st.secrets["GEMINI_API_KEY"]
    
    try:
        genai.configure(api_key=api_key)
        # 使用最新的 Gemini 2.5 Flash 模型
        model = genai.GenerativeModel('gemini-2.5-flash')
        
        system_prompt = """
        你是一位擁有 20 年經驗的企業流程優化專家。請將用戶輸入的雜亂資訊，整理成一份專業、結構化、可直接執行的標準作業程序 (SOP)。
        
        【輸出格式要求】：
        1. 使用標準 Markdown 格式。
        2. 標題層級清晰 (## 目標, ### 步驟)。
        3. 必須包含以下區塊：
           - **目標 (Objective)**：一句話說明此流程目的。
           - **適用對象 (Scope)**：誰該執行此流程。
           - **前置準備 (Prerequisites)**：需要的工具、權限或材料。
           - **詳細執行步驟 (Procedure)**：條列式步驟，關鍵動作請加粗。
           - **風險與注意事項 (Risks & Notes)**：可能的雷區。
        4. 語氣專業、精煉，避免廢話。
        """
        
        with st.spinner("🤖 AI 正在分析流程並撰寫文檔... (約需 10-20 秒)"):
            response = model.generate_content(f"{system_prompt}\n\n【用戶輸入內容】：\n{raw_text}")
            return response.text
            
    except Exception as e:
        st.error(f"生成失敗，請稍後再試。錯誤訊息：{e}")
        return None

def create_docx(text):
    """將 Markdown 文字轉為 Word 檔"""
    doc = Document()
    doc.add_heading('標準作業程序 (SOP)', 0)
    
    for line in text.split('\n'):
        line = line.strip()
        if not line: continue
        
        if line.startswith('## '):
            doc.add_heading(line.replace('## ', ''), level=1)
        elif line.startswith('### '):
            doc.add_heading(line.replace('### ', ''), level=2)
        elif line.startswith('* ') or line.startswith('- '):
            doc.add_paragraph(line[2:], style='List Bullet')
        elif line[0].isdigit() and line[1] == '.':
            doc.add_paragraph(line, style='List Number')
        else:
            doc.add_paragraph(line)
            
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- 3. 主程式邏輯 ---

# 🔒 第一關：門禁檢查
if not check_password():
    st.title(PAGE_TITLE)
    st.info("👈 請在左側欄位輸入密碼以開始使用。")
    st.stop() # 沒密碼就停在這裡

# 🔓 第二關：進入主畫面
st.title(PAGE_TITLE)
st.success(f"✅ VIP 驗證成功！剩餘生成額度：{MAX_USAGE_PER_SESSION - st.session_state['usage_count']} 次")

col1, col2 = st.columns([1, 1])

with col1:
    st.subheader("1. 輸入資料")
    user_input = st.text_area("請貼上會議記錄、語音轉文字稿或雜亂筆記...", height=400)
    
    # 生成按鈕
    if st.button("🚀 生成專業 SOP", type="primary", use_container_width=True):
        if user_input:
            # 檢查額度
            if check_usage_limit():
                result = generate_sop(user_input)
                if result:
                    st.session_state['result'] = result
                    st.session_state['usage_count'] += 1 # 扣除額度
                    st.rerun() # 重新整理以更新顯示的剩餘次數
        else:
            st.warning("請先輸入內容！")

with col2:
    st.subheader("2. 生成結果")
    if st.session_state['result']:
        # 顯示預覽
        st.markdown(st.session_state['result'])
        st.divider()
        
        # 製作 Word 檔
        docx_file = create_docx(st.session_state['result'])
        
        st.download_button(
            label="📥 下載 Word 檔案 (.docx)",
            data=docx_file,
            file_name="SOP_Output.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            type="primary",
            use_container_width=True
        )
    else:
        st.info("👈 在左側輸入內容並點擊生成，結果將顯示於此。")

